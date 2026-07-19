import os
import sys
from pathlib import Path
import pandas as pd
from PIL import Image, ImageDraw
import docx
import fitz

# Ensure the parent directory is in sys.path so we can import modules
sys.path.append(str(Path(__file__).resolve().parent.parent))

from parsers.pdf_parser import parse_pdf
from parsers.image_parser import parse_image
from parsers.docx_parser import parse_docx
from parsers.excel_parser import parse_excel

TEST_DIR = Path(__file__).resolve().parent / "sample_files"

def setup_test_files():
    TEST_DIR.mkdir(exist_ok=True)
    
    # 1. Create a dummy PDF
    pdf_path = TEST_DIR / "test.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Test PDF Document\nDigital Page")
    
    # Add a "scanned" page (just very little text to trigger threshold)
    page2 = doc.new_page()
    page2.insert_text((50, 50), "hi") 
    doc.save(str(pdf_path))
    doc.close()
    
    # 2. Create a dummy image
    img_path = TEST_DIR / "test.png"
    img = Image.new('RGB', (200, 100), color = (255, 255, 255))
    d = ImageDraw.Draw(img)
    d.text((10,10), "Test Image OCR", fill=(0,0,0))
    img.save(str(img_path))
    
    # 3. Create a dummy DOCX
    docx_path = TEST_DIR / "test.docx"
    doc = docx.Document()
    doc.add_heading('Main Heading', level=1)
    doc.add_paragraph('This is a test paragraph for DOCX parsing.')
    doc.add_heading('Sub Heading', level=2)
    doc.add_paragraph('Another paragraph.')
    doc.save(str(docx_path))
    
    # 4. Create a dummy XLSX
    xlsx_path = TEST_DIR / "test.xlsx"
    df1 = pd.DataFrame({"Name": ["Alice", "Bob"], "Age": [25, 30]})
    df2 = pd.DataFrame({"Item": ["Apple", "Banana"], "Price": [1.2, 0.5]})
    
    with pd.ExcelWriter(str(xlsx_path)) as writer:
        df1.to_excel(writer, sheet_name="People", index=False)
        df2.to_excel(writer, sheet_name="Products", index=False)
        
    return pdf_path, img_path, docx_path, xlsx_path

def main():
    print("=== Generating test files ===")
    pdf_path, img_path, docx_path, xlsx_path = setup_test_files()
    
    print("\n=== Testing PDF Parser ===")
    try:
        pdf_res = parse_pdf(pdf_path)
        print(f"Result: {pdf_res}")
    except Exception as e:
        print(f"PDF Parser Failed: {e}")

    print("\n=== Testing Image Parser (OCR) ===")
    try:
        img_res = parse_image(img_path)
        print(f"Result: {img_res}")
    except Exception as e:
        print(f"Image Parser Failed: {e}")
        
    print("\n=== Testing DOCX Parser ===")
    try:
        docx_res = parse_docx(docx_path)
        print(f"Result: {docx_res}")
    except Exception as e:
        print(f"DOCX Parser Failed: {e}")
        
    print("\n=== Testing Excel Parser ===")
    try:
        xlsx_res = parse_excel(xlsx_path)
        print(f"Result: {xlsx_res}")
    except Exception as e:
        print(f"Excel Parser Failed: {e}")

if __name__ == "__main__":
    main()
